# 利用python编辑word测试

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Length

from docx.document import Document as Doc
# from sqlalchemy import true

from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建代表Word文档的Doc对象
document = Document()  # type: Doc
section = document.sections[0]
# 添加页眉
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = "python学习-操纵word"
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # 对齐方式：居中
# 添加页脚
footer = section.footer
paragraph = footer.paragraphs[0]
paragraph.text = "仅供学习参考"
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # 对齐方式：居中
# 添加大标题
document.add_heading('Python编辑word测试', 0)
# 添加段落
p = document.add_paragraph('Python是一门非常流行的编程语言，它')
run = p.add_run('简单')
run.bold = True
run.font.size = Pt(18)
p.add_run('而且')
run = p.add_run('优雅')
run.font.size = Pt(18)
run.underline = True
p.add_run('。')

p1 = document.add_paragraph()
p1_format = p1.paragraph_format
p1_format.alignment = WD_ALIGN_PARAGRAPH.LEFT # 对齐方式：左对齐
text1 = p1.add_run("数据审查")
text1.font.size = Pt(15)    #小三字号
text1.bold = True
text1.font.name = 'Times New Roman'           # 控制是西文时的字体
text1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p2 = document.add_paragraph()
#p2.line_spacing_rule = WD_LINE_SPACING.MULTIPLE #多倍行距
p2_format = p2.paragraph_format
p2_format.first_line_indent = Inches(0.35)
p2_format.line_spacing = 1.5    # 1.5倍行间距
p2_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY    # 对齐方式：两端对齐
text2 = p2.add_run("客户在系统中提交中试申请、技术资料、入库通知单、源程序等，安全质量部需对客户发出的申请、测试资料等进行审查，审查不通过的需打回给客户并说明原因；审查通过的则将审查资料存入系统，并生成数据审查记录。")
text2.font.size = Pt(12)        #小四字号
text2.font.name = 'Times New Roman'           # 控制是西文时的字体
text2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')



# # 添加一级标题
# document.add_heading('Heading, level 1', level=1)
# # 添加带样式的段落
# #document.add_paragraph('Intense quote', style='Intense Quote')
# # 添加无序列表
# document.add_paragraph('first item in unordered list', style='List Bullet')
# document.add_paragraph('second item in ordered list', style='List Bullet')

# # 添加有序列表
# document.add_paragraph('first item in ordered list', style='List Number')
# document.add_paragraph('second item in ordered list', style='List Number')

# # 添加图片（注意路径和图片必须要存在）
# #document.add_picture('resources/guido.jpg', width=Cm(5.2))

# 添加分节符
# document.add_section()

records = (
    ('骆昊', '男', '1995-5-5'),
    ('孙美丽', '女', '1992-2-2'),
    ('刘三姐','女','1999-5-4')
)
# 添加表格
table = document.add_table(rows=1, cols=3)
table.style = 'TableGrid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '姓名'
hdr_cells[1].text = '性别'
hdr_cells[2].text = '出生日期'

# 表格首行加粗
for i in range(3):
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

# 为表格添加行
for item in records:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]




# 添加分页符
#document.add_page_break()

# 保存文档
document.save('demo.docx')