# coding:utf-8
# 读取word文件测试

import os
from docx import Document


path = os.path.join(os.getcwd(), 'demo.docx')
print("文件的路径为：", path)     # 调试路径

doc = Document(path)

for p in doc.paragraphs:
    print(p.text)

for t in doc.tables:            # for 循环获取表格对象
    for row in t.rows:          # 获取每一行
        row_str = []
        for cell in row.cells:    # 获取每一行单独的小表格,然后将其内容拼接起来;拼接完成之后再第二个for循环中打印出来
            row_str.append(cell.text)
        print(row_str)
        

